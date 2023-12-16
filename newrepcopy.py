import tkinter as tk
from tkinter import filedialog
import re
from docx import Document

# Declare val1 and val2 as global variables
val1 = None
val2 = None
line_number =None
page_number = None
def check_vals_in_second_file(val1, val2, second_doc):
    #val1_present = any(val1 in paragraph.text for paragraph in second_doc.paragraphs)
    val2_present = any(val2 in paragraph.text for paragraph in second_doc.paragraphs)
    print(val2_present)
    return val2_present

def check_vals_in_second_file(val1, val2, second_doc):
    val2_present = any(val2 in paragraph.text for paragraph in second_doc.paragraphs)
    return val2_present

def process_document(doc_content, second_doc):
    global val1, val2  # Update val1 and val2 in the global scope
    # Flag to indicate when to start processing content
    start_processing = False
    b_value = None
    current_change = None  # To store information about the current change

    for paragraph in doc_content.paragraphs:
        # Check for the starting point
        change_match = re.match(r'Change (\d+):', paragraph.text)
        if change_match:
            # If a new change section is found, print the values from the previous section (if b=1)
            if current_change is not None and b_value == 1:
                print("Change:", current_change)
                print("Val1:", val1)
                print("Val2:", val2)
                if check_vals_in_second_file(val1, val2, second_doc):
                    print("YES")
                    # line_number, page_number = get_line_and_page(paragraph.text)
                    print("Line:", line_number)
                    print("Page:", page_number)
                    update_paragraph_in_second_doc(second_doc ,line_number, page_number, val1 )
                    #save_document(second_doc, "modified_second_doc.docx")
                else:
                    print("NO")

            # Reset values for the new change section
            val1 = None
            val2 = None
            b_value = None
            start_processing = True
            current_change = change_match.group(1)
            continue

        # Process content only if the flag is True
        if start_processing:
            match_b = re.search(r'b\s*=\s*(\d+)', paragraph.text)
            if match_b:
                b_value = int(match_b.group(1))
            match_line_page = re.search(r'\(Line (\d+), Page (\d+)\):', paragraph.text)
            if match_line_page:
                line_number = match_line_page.group(1)
                page_number = match_line_page.group(2)
                # print("Line:", line_number)
                # print("Page:", page_number)

            match = re.search(r':\s*(.+)', paragraph.text)
            if match:
                val1_candidate = match.group(1).lstrip()  # Strip leading spaces
                if val1 is None and val1_candidate:  # Only update val1 if it's not set and not an empty string
                    val1 = val1_candidate
                else:
                    val2 = match.group(1).strip()  # Strip leading and trailing spaces

    # Print values for the last change section (if b=1)
    if current_change is not None and b_value == 1:
        print("Change:", current_change)
        print("Val1:", val1)
        print("Val2:", val2)
        print("Line:", line_number)
        print("Page:", page_number)
        update_paragraph_in_second_doc(second_doc ,line_number, page_number,val1 )
        
    save_document(second_doc, "modified_second_doc.docx")


def save_document(doc, filename):
    doc.save(filename)

def update_paragraph_in_second_doc(second_doc, line_number, page_number, new_content):
    # Modify line_number and page_number as needed
    modified_line_number = line_number.replace('-', '_')  # Replace '-' with '_'
    modified_page_number = page_number.replace('/', '_')  # Replace '/' with '_'

    # Find the specified paragraph in the second document based on modified Line and Page numbers
    for i, paragraph in enumerate(second_doc.paragraphs):
        # print(paragraph.text)
        if i == (int(line_number)-1) : 
            second_doc.paragraphs[i].clear()
            print("______________________________"+new_content)
            second_doc.paragraphs[i].add_run(f"{new_content}")
        

def get_line_and_page(text):
    match_line_page = re.search(r'\(Line (\d+), Page (\d+)\):', text)
    if match_line_page:
        line_number = match_line_page.group(1)
        page_number = match_line_page.group(2)
        return line_number, page_number
    return None, None




def upload_file():
    global val1, val2  # Ensure we are using the global val1 and val2 variables
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    
    try:
        doc = Document(file_path)
        

        # Ask the user for a second Word file
        second_file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        second_doc = Document(second_file_path)
        process_document(doc , second_doc)
        # Check if val1 and val2 are present in the second document
        
        

    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found.")
    except Exception as e:
        print(f"Error: {e}")

# Create the main application window
app = tk.Tk()
app.title("File Upload Example")

# Create and pack the upload button
upload_button = tk.Button(app, text="Upload Word File", command=upload_file)
upload_button.pack(pady=20)

# Run the main loop
app.mainloop()
