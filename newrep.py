import tkinter as tk
from tkinter import filedialog
import re
from docx import Document

# Declare val1 and val2 as global variables
val1 = None
val2 = None


def check_vals_in_second_file(val1, val2, second_doc):
    #val1_present = any(val1 in paragraph.text for paragraph in second_doc.paragraphs)
    val2_present = any(val2 in paragraph.text for paragraph in second_doc.paragraphs)
    print(val2_present)
    return val2_present

def process_document(doc_content , second_doc):
    global val1, val2  # Update val1 and val2 in the global scope
    # Flag to indicate when to start processing content
    start_processing = False
    b_value = None

    for paragraph in doc_content.paragraphs:
        # Check for the starting point
        change_match = re.match(r'Change (\d+):', paragraph.text)
        if change_match:
            # If a new change section is found, print the values from the previous section (if b=1)
            if val1 is not None and val2 is not None and b_value == 1:
                print("Val1:", val1)
                print("Val2:", val2)
                if  check_vals_in_second_file(val1, val2, second_doc):
                    print("YES")
                else:
                    print("NO")
            # Reset values for the new change section
            val1 = None
            val2 = None
            b_value = None
            start_processing = True
            continue

        # Process content only if the flag is True
        if start_processing:
            match_b = re.search(r'b\s*=\s*(\d+)', paragraph.text)
            if match_b:
                b_value = int(match_b.group(1))
            match = re.search(r':\s*(.+)', paragraph.text)
            if match:
                val1_candidate = match.group(1).lstrip()  # Strip leading spaces
                if val1 is None and val1_candidate:  # Only update val1 if it's not set and not an empty string
                    val1 = val1_candidate
                else:
                    val2 = match.group(1).strip()  # Strip leading and trailing spaces

    # Print values for the last change section (if b=1)
    if val1 is not None and val2 is not None and b_value == 1:
        print("Val1:", val1)
        print("Val2:", val2)
       



def upload_file():
    global val1, val2  # Ensure we are using the global val1 and val2 variables
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    
    try:
        doc = Document(file_path)
        

        # Ask the user for a second Word file
        second_file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        second_doc = Document(second_file_path)
        process_document(doc , second_doc)
        # Check if val1 and val2 are present in the second document
        
        

    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found.")
    except Exception as e:
        print(f"Error: {e}")

# Create the main application window
app = tk.Tk()
app.title("File Upload Example")

# Create and pack the upload button
upload_button = tk.Button(app, text="Upload Word File", command=upload_file)
upload_button.pack(pady=20)

# Run the main loop
app.mainloop()
